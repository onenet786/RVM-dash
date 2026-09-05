import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def markdown_to_docx(md_path, docx_path, title_text, subtitle_text):
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Base Styles with Arial font
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(51, 65, 85) # Slate 700

    # Header / Title Banner Table
    title_table = doc.add_table(rows=1, cols=1)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_cell = title_table.cell(0, 0)
    set_cell_background(title_cell, "0F172A") # Dark Slate Navy
    set_cell_margins(title_cell, top=280, bottom=280, left=350, right=350)
    
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_badge = p.add_run("RVM DESKTOP APPLICATION • OFFICIAL SYSTEM DOCUMENTATION\n")
    run_badge.font.name = 'Arial'
    run_badge.font.size = Pt(9.5)
    run_badge.font.bold = True
    run_badge.font.color.rgb = RGBColor(56, 189, 248) # Cyan Accent
    
    run_title = p.add_run(title_text + "\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)
    
    run_sub = p.add_run(subtitle_text)
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(203, 213, 225) # Light Slate

    doc.add_paragraph() # Spacing

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_data = []

    for raw_line in lines:
        line = raw_line.rstrip('\r\n')

        # Check Code block delimiters
        if line.strip().startswith('```'):
            if in_code_block:
                # Flush code block
                in_code_block = False
                block_text = "\n".join(code_lines)
                code_lines = []
                
                # Check if it's a screenshot placeholder or code
                is_placeholder = "SCREENSHOT PLACEHOLDER" in block_text
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                c = tbl.cell(0, 0)
                
                if is_placeholder:
                    set_cell_background(c, "F1F5F9") # Slate Light
                    set_cell_margins(c, top=140, bottom=140, left=200, right=200)
                    cp = c.paragraphs[0]
                    r = cp.add_run("📸 " + block_text.replace("+", "").replace("|", "").strip())
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(15, 118, 110) # Teal
                else:
                    set_cell_background(c, "1E293B") # Dark slate
                    set_cell_margins(c, top=140, bottom=140, left=200, right=200)
                    cp = c.paragraphs[0]
                    r = cp.add_run(block_text)
                    r.font.name = 'Consolas'
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = RGBColor(241, 245, 249)
                doc.add_paragraph()
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Check Table
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Table row
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            # Check if separator row like |:---|:---|
            if all(set(c).issubset({'-', ':', ' '}) for c in cells):
                continue
            table_data.append(cells)
            in_table = True
            continue
        else:
            if in_table:
                # Flush table
                in_table = False
                if table_data:
                    num_rows = len(table_data)
                    num_cols = max(len(r) for r in table_data)
                    tbl = doc.add_table(rows=num_rows, cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for r_idx, row in enumerate(table_data):
                        for c_idx, val in enumerate(row):
                            if c_idx < num_cols:
                                cell = tbl.cell(r_idx, c_idx)
                                cell.text = val
                                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                                p = cell.paragraphs[0]
                                p.runs[0].font.name = 'Arial'
                                p.runs[0].font.size = Pt(9.5)
                                if r_idx == 0:
                                    set_cell_background(cell, "0284C7") # Ocean Blue Header
                                    p.runs[0].font.bold = True
                                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                                else:
                                    bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                                    set_cell_background(cell, bg)
                                    p.runs[0].font.color.rgb = RGBColor(30, 41, 59)
                    doc.add_paragraph()
                table_data = []

        # Empty lines
        if not line.strip():
            continue

        # Headings
        if line.startswith('# '):
            continue # Already handled in title banner
        elif line.startswith('## '):
            h_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(14, 116, 144) # Deep Teal
            # Bottom border line
            continue
        elif line.startswith('### '):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
            continue
        elif line.startswith('#### '):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(3, 105, 161) # Sky 700
            continue

        # Bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, bullet_text)
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if num_match:
            item_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, item_text)
            continue

        # Horizontal rule
        if line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run("____________________________________________________________________")
            r.font.name = 'Arial'
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(226, 232, 240)
            continue

        # Normal Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        format_inline_markdown(p, line)

    doc.save(docx_path)
    print(f"[SUCCESS] Built DOCX: {docx_path}")

def format_inline_markdown(paragraph, text):
    # Splits by bold `**text**` and code `` `code` ``
    tokens = re.split(r'(\*\*.*?\*\*|\`.*?\`|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            r = paragraph.add_run(token[2:-2])
            r.font.name = 'Arial'
            r.font.bold = True
            r.font.color.rgb = RGBColor(15, 23, 42)
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Arial'
            r.font.italic = True
            r.font.color.rgb = RGBColor(71, 85, 105)
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(180, 83, 9) # Amber 700
        else:
            r = paragraph.add_run(token)
            r.font.name = 'Arial'
            r.font.color.rgb = RGBColor(51, 65, 85)

if __name__ == '__main__':
    base_dir = r"d:\GIT-HUB\RVM-dash\docs\rvm_desktop_app_docs"
    
    # 1. User Guide
    ug_md = os.path.join(base_dir, "RVMDesktopApp_User_Guide.md")
    ug_docx = os.path.join(base_dir, "RVMDesktopApp_User_Guide.docx")
    markdown_to_docx(ug_md, ug_docx, 
                     "RVMDesktopApp User Guide", 
                     "Step-by-Step Module Manual, Role Guides, Hardware Diagnostics & FAQs\nLanguage: Roman Urdu + English Mix | Framework: .NET 8.0 WPF / Arduino Uno")
                     
    # 2. Logic Explanation
    le_md = os.path.join(base_dir, "RVMDesktopApp_Logic_Explanation.md")
    le_docx = os.path.join(base_dir, "RVMDesktopApp_Logic_Explanation.docx")
    markdown_to_docx(le_md, le_docx, 
                     "RVMDesktopApp Architecture & Logic", 
                     "End-to-End Data Pipeline, Module Dependencies, 'Why' Matrix & Business Rules\nLanguage: Roman Urdu + English Mix | Engineering Reference Manual")
